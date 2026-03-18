import os
import PyPDF2
from docx import Document
from typing import List, Dict
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class DocumentLoader:
    """文档加载器，支持PDF和DOCX格式"""
    
    @staticmethod
    def load_pdf(file_path: str) -> Dict[str, any]:
        """加载PDF文档"""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = []
                metadata = {
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'pages': len(pdf_reader.pages)
                }
                
                for page_num, page in enumerate(pdf_reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        text_content.append({
                            'page': page_num + 1,
                            'content': text
                        })
                
                full_text = ' '.join([page['content'] for page in text_content])
                
                return {
                    'type': 'pdf',
                    'metadata': metadata,
                    'pages': text_content,
                    'full_text': full_text,
                    'success': True
                }
        except Exception as e:
            logger.error(f"Error loading PDF: {str(e)}")
            return {
                'type': 'pdf',
                'error': str(e),
                'success': False
            }
    
    @staticmethod
    def load_docx(file_path: str) -> Dict[str, any]:
        """加载DOCX文档"""
        try:
            doc = Document(file_path)
            text_content = []
            
            # 提取段落
            for para_num, paragraph in enumerate(doc.paragraphs):
                if paragraph.text.strip():
                    text_content.append({
                        'paragraph': para_num + 1,
                        'content': paragraph.text
                    })
            
            # 提取表格
            tables_content = []
            for table_num, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                tables_content.append({
                    'table': table_num + 1,
                    'data': table_data
                })
            
            full_text = ' '.join([para['content'] for para in text_content])
            
            metadata = {
                'title': doc.core_properties.title or '',
                'author': doc.core_properties.author or '',
                'paragraphs': len(text_content),
                'tables': len(tables_content)
            }
            
            return {
                'type': 'docx',
                'metadata': metadata,
                'paragraphs': text_content,
                'tables': tables_content,
                'full_text': full_text,
                'success': True
            }
        except Exception as e:
            logger.error(f"Error loading DOCX: {str(e)}")
            return {
                'type': 'docx',
                'error': str(e),
                'success': False
            }
    
    @staticmethod
    def load_document(file_path: str) -> Dict[str, any]:
        """根据文件类型自动加载文档"""
        file_extension = os.path.splitext(file_path)[1].lower()
        
        if file_extension == '.pdf':
            return DocumentLoader.load_pdf(file_path)
        elif file_extension == '.docx':
            return DocumentLoader.load_docx(file_path)
        else:
            return {
                'error': f'Unsupported file type: {file_extension}',
                'success': False
            }
    
    @staticmethod
    def extract_sections(text: str) -> Dict[str, List[str]]:
        """从文档中提取主要章节"""
        sections = {
            'abstract': [],
            'introduction': [],
            'methodology': [],
            'results': [],
            'conclusion': [],
            'references': [],
            'other': []
        }
        
        # 简单的章节检测逻辑
        lines = text.split('\n')
        current_section = 'other'
        
        for line in lines:
            line_lower = line.lower().strip()
            
            if any(keyword in line_lower for keyword in ['abstract', '摘要']):
                current_section = 'abstract'
            elif any(keyword in line_lower for keyword in ['introduction', '引言', '介绍']):
                current_section = 'introduction'
            elif any(keyword in line_lower for keyword in ['method', 'methodology', '方法', '方法论']):
                current_section = 'methodology'
            elif any(keyword in line_lower for keyword in ['result', '结果', '实验结果']):
                current_section = 'results'
            elif any(keyword in line_lower for keyword in ['conclusion', '结论', '总结']):
                current_section = 'conclusion'
            elif any(keyword in line_lower for keyword in ['reference', '参考文献', '引用']):
                current_section = 'references'
            
            if line.strip():
                sections[current_section].append(line)
        
        return {k: '\n'.join(v) for k, v in sections.items()}
